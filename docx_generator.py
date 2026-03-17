from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x0f, 0x20, 0x44)
MID_NAVY   = RGBColor(0x1a, 0x3a, 0x6b)
ACCENT     = RGBColor(0x25, 0x63, 0xeb)
TEXT_DARK  = RGBColor(0x1e, 0x29, 0x3b)
TEXT_MID   = RGBColor(0x47, 0x55, 0x69)
TEXT_LIGHT = RGBColor(0x94, 0xa3, 0xb8)
ROW_ALT    = RGBColor(0xf1, 0xf5, 0xf9)
ROW_HEAD   = RGBColor(0x0f, 0x20, 0x44)
WHITE      = RGBColor(0xff, 0xff, 0xff)


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Set individual cell borders. Pass top/bottom/left/right dicts."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, props in kwargs.items():
        border = OxmlElement(f"w:{side}")
        for k, v in props.items():
            border.set(qn(f"w:{k}"), str(v))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_run(para, text, bold=False, italic=False,
             color: RGBColor = None, size_pt: int = None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _para(doc, text="", bold=False, italic=False,
          color: RGBColor = None, size_pt: int = 10,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=4, font="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _add_run(p, text, bold=bold, italic=italic,
                 color=color, size_pt=size_pt, font=font)
    return p


def _add_bottom_border(para, hex_color="2563eb", size="12"):
    """Add a bottom border to a paragraph (used for section headings)."""
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Cover page ────────────────────────────────────────────────────────────────
def _cover_page(doc, structured_data):
    today = date.today().strftime("%B %d, %Y")
    best_tx    = structured_data.get("best_treatment", "—")
    conf       = structured_data.get("confidence_score", "—")
    warning    = structured_data.get("warning", "—")
    conclusion = structured_data.get("conclusion", "—").upper()

    # Large dark cover block via a 1-cell table
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover_tbl.cell(0, 0)
    _set_cell_bg(cell, "0f2044")

    # Title
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(40)
    tp.paragraph_format.space_after  = Pt(6)
    _add_run(tp, "CLINICAL STUDY REPORT",
             bold=True, color=WHITE, size_pt=28, font="Calibri")

    # Subtitle
    sp = cell.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(4)
    _add_run(sp, "Prepared by TrialDocAI  ·  Powered by Groq LLaMA 3.3 70B",
             color=RGBColor(0x93, 0xc5, 0xfd), size_pt=10, font="Calibri")

    # Meta lines
    for line in [f"Date: {today}", "Sponsor: Demo Sponsor Inc.",
                 "Document Status: DRAFT — FOR REVIEW"]:
        mp = cell.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(3)
        _add_run(mp, line, color=RGBColor(0xcb, 0xd5, 0xe1), size_pt=9, font="Calibri")

    # Spacer
    sp2 = cell.add_paragraph()
    sp2.paragraph_format.space_after = Pt(20)

    # Remove the default empty first paragraph in cell
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    # ── Metrics bar (separate table below cover) ──────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    metrics = [
        ("BEST TREATMENT", str(best_tx)),
        ("CONCLUSION",     str(conclusion)),
        ("EFFECT",         str(warning)),
        ("CONFIDENCE",     str(conf)),
    ]
    m_tbl = doc.add_table(rows=2, cols=4)
    m_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_i, (lbl, val) in enumerate(metrics):
        lbl_cell = m_tbl.cell(0, col_i)
        val_cell = m_tbl.cell(1, col_i)
        _set_cell_bg(lbl_cell, "1a3a6b")
        _set_cell_bg(val_cell, "1a3a6b")

        lp = lbl_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after  = Pt(2)
        _add_run(lp, lbl, bold=True,
                 color=RGBColor(0x94, 0xa3, 0xb8), size_pt=7.5, font="Calibri")

        vp = val_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after  = Pt(8)
        _add_run(vp, val, bold=True, color=WHITE, size_pt=13, font="Calibri")

    doc.add_page_break()


# ── Section heading ───────────────────────────────────────────────────────────
def _section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    _add_run(p, f"{number}.  ", bold=True,
             color=ACCENT, size_pt=13, font="Calibri")
    _add_run(p, title, bold=True,
             color=DARK_NAVY, size_pt=13, font="Calibri")
    _add_bottom_border(p, hex_color="2563eb", size="8")
    return p


# ── Results table ─────────────────────────────────────────────────────────────
def _results_table(doc, structured_data):
    results = structured_data.get("results", [])
    if not results:
        return

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after  = Pt(4)
    _add_run(cap, "Table 1. ", bold=True, color=MID_NAVY, size_pt=9, font="Calibri")
    _add_run(cap, "Primary Endpoint Results — Blood Pressure Reduction (mmHg)",
             italic=True, color=TEXT_MID, size_pt=9, font="Calibri")

    headers = ["Treatment Group", "N", "Mean Reduction (mmHg)", "Std Dev"]
    tbl = doc.add_table(rows=1 + len(results), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(2.4), Inches(0.8), Inches(2.0), Inches(1.0)]

    # Header row
    hdr_row = tbl.rows[0]
    for i, (cell, w, h) in enumerate(zip(hdr_row.cells, col_widths, headers)):
        cell.width = w
        _set_cell_bg(cell, "0f2044")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        _add_run(p, h, bold=True, color=WHITE, size_pt=9, font="Calibri")

    # Data rows
    for row_i, r in enumerate(results):
        row = tbl.rows[row_i + 1]
        row_data = [r["treatment"], str(r["sample_size"]),
                    str(r["mean"]), str(r["std_dev"])]
        bg = "f1f5f9" if row_i % 2 else "ffffff"
        for col_i, (cell, w, val) in enumerate(zip(row.cells, col_widths, row_data)):
            cell.width = w
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            _add_run(p, val, color=TEXT_DARK, size_pt=9.5, font="Calibri")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    _add_run(note, "Note: Values represent mean change from baseline ± standard deviation.",
             italic=True, color=TEXT_MID, size_pt=8, font="Calibri")


# ── Pairwise stats table ──────────────────────────────────────────────────────
def _pairwise_table(doc, structured_data):
    pairwise = structured_data.get("pairwise_stats", [])
    if not pairwise:
        return

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after  = Pt(4)
    _add_run(cap, "Table 2. ", bold=True, color=MID_NAVY, size_pt=9, font="Calibri")
    _add_run(cap, "Pairwise Statistical Comparisons vs Placebo (Welch's t-test)",
             italic=True, color=TEXT_MID, size_pt=9, font="Calibri")

    headers = ["Treatment", "t-statistic", "p-value", "Significance", "Cohen's d", "Effect"]
    col_widths = [Inches(1.0), Inches(0.95), Inches(0.8), Inches(2.0), Inches(0.8), Inches(0.75)]

    tbl = doc.add_table(rows=1 + len(pairwise), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl.rows[0]
    for cell, w, h in zip(hdr_row.cells, col_widths, headers):
        cell.width = w
        _set_cell_bg(cell, "1a3a6b")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        _add_run(p, h, bold=True, color=WHITE, size_pt=8.5, font="Calibri")

    for row_i, ps in enumerate(pairwise):
        row = tbl.rows[row_i + 1]
        row_data = [ps["treatment"], str(ps["t_statistic"]), str(ps["p_value"]),
                    ps["significance"], str(ps["cohens_d"]), ps["effect_size_label"]]
        bg = "f1f5f9" if row_i % 2 else "ffffff"
        for col_i, (cell, w, val) in enumerate(zip(row.cells, col_widths, row_data)):
            cell.width = w
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            _add_run(p, val, color=TEXT_DARK, size_pt=8.5, font="Calibri")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    _add_run(note,
             "Welch's t-test used (unequal variance assumed). "
             "Cohen's d: Small <0.5, Medium 0.5–0.8, Large >0.8.",
             italic=True, color=TEXT_MID, size_pt=8, font="Calibri")


# ── Main function ─────────────────────────────────────────────────────────────
def create_docx_bytes(report_text: str, structured_data: dict) -> BytesIO:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover page ────────────────────────────────────────────────────────────
    _cover_page(doc, structured_data)

    # ── Stats tables section header ───────────────────────────────────────────
    lbl = doc.add_paragraph()
    lbl.paragraph_format.space_after = Pt(2)
    _add_run(lbl, "STATISTICAL RESULTS SUMMARY",
             bold=True, color=ACCENT, size_pt=8, font="Calibri")
    _add_bottom_border(lbl, hex_color="2563eb", size="12")

    _results_table(doc, structured_data)
    _pairwise_table(doc, structured_data)

    # ── CSR narrative ─────────────────────────────────────────────────────────
    lbl2 = doc.add_paragraph()
    lbl2.paragraph_format.space_before = Pt(14)
    lbl2.paragraph_format.space_after  = Pt(2)
    _add_run(lbl2, "CLINICAL STUDY REPORT",
             bold=True, color=ACCENT, size_pt=8, font="Calibri")
    _add_bottom_border(lbl2, hex_color="2563eb", size="12")

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # Markdown bold heading: **Title**
        if line.startswith("**") and line.endswith("**"):
            clean = line.replace("**", "")
            match = re.match(r"^(\d+)\.\s+(.+)$", clean)
            if match:
                _section_heading(doc, match.group(1), match.group(2))
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
                _add_run(p, clean, bold=True,
                         color=DARK_NAVY, size_pt=13, font="Calibri")
                _add_bottom_border(p)
            continue

        # Numbered heading without bold markers: "1. Title Page"
        match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if match and len(line) < 80:
            _section_heading(doc, match.group(1), match.group(2))
            continue

        # Markdown pipe table rows — skip
        if line.startswith("|"):
            continue

        # Normal body paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        _add_run(p, line, color=TEXT_DARK, size_pt=10, font="Calibri")

    # ── Footer disclaimer ─────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(disc,
             "This document was generated by TrialDocAI for demonstration purposes. "
             "Not for regulatory submission without medical writer review.",
             italic=True, color=TEXT_LIGHT, size_pt=8, font="Calibri")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer